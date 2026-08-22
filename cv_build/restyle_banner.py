# -*- coding: utf-8 -*-
"""Produce the 'banner' Word file from the Gulim one — same content, new skin.

Everything is read from SeunghanLee_CV.docx, so the text can never drift:
only fonts, sizes, colours, bullets, borders and the page furniture change.
Lato is embedded in the file so it looks right on a PC without Lato installed.
"""
import os
import shutil
import struct
import uuid
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = "/nfsdata/home/seunghan.lee/web/SeunghanLee_CV.docx"
OUT = "/nfsdata/home/seunghan.lee/web/cv_candidates/SeunghanLee_CV_banner.docx"
DL = os.path.join(HERE, "fonts_dl")

FONT = "Lato"
SCALE = 0.85
ACCENT = RGBColor(0x0F, 0x3D, 0x57)
ACCENT2 = RGBColor(0xC4, 0x76, 0x1A)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
HEAD_PT = 14.5

COLOR_MAP = {
    "1F497D": ACCENT, "FF9900": ACCENT2, "212121": MUTED,
    "333333": MUTED, "222222": MUTED, "0000FF": ACCENT,
}
BULLET_MAP = {"•": ("•", 8.5), "o": ("◦", 8.0), "§": ("–", 8.5)}

doc = Document(SRC)


# ---------------------------------------------------------------- helpers
def iter_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in iter_paragraphs(cell):
                    yield p


def hyperlink_runs(par):
    ids = set()
    for hl in par._p.findall(qn("w:hyperlink")):
        for r in hl.findall(qn("w:r")):
            ids.add(id(r))
    return ids


def all_runs(par):
    from docx.text.run import Run

    for child in par._p.iterchildren():
        if child.tag == qn("w:r"):
            yield Run(child, par)
        elif child.tag == qn("w:hyperlink"):
            for r in child.findall(qn("w:r")):
                yield Run(r, par)


def set_font(run, name=FONT):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)


def para_border(par):
    pPr = par._p.pPr
    if pPr is None:
        return None
    bdr = pPr.find(qn("w:pBdr"))
    return None if bdr is None else bdr.find(qn("w:bottom"))


def set_border(bottom, color, sz):
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)


# ---------------------------------------------------------------- restyle
n_head = 0
for par in iter_paragraphs(doc):
    pf = par.paragraph_format
    bottom = para_border(par)
    is_heading = bottom is not None
    if is_heading:
        set_border(bottom, "C4761A", 14)     # thin amber rule instead of the black bar
        n_head += 1
    if pf.left_indent is not None:
        pf.left_indent = Cm(pf.left_indent.cm * 0.92)
    if pf.first_line_indent is not None:
        pf.first_line_indent = Cm(pf.first_line_indent.cm * 0.92)
    if pf.space_before is not None:
        pf.space_before = Pt(round(pf.space_before.pt * 1.1, 1))

    hl = hyperlink_runs(par)
    for run in all_runs(par):
        set_font(run)
        txt = run.text.strip()
        if txt in BULLET_MAP:                # level bullets
            ch, sz = BULLET_MAP[txt]
            run.text = run.text.replace(txt, ch)
            run.font.size = Pt(sz)
            run.font.color.rgb = ACCENT
            continue
        if run.font.size is not None:
            pt = HEAD_PT if is_heading else round(run.font.size.pt * SCALE * 2) / 2
            run.font.size = Pt(pt)
        col = run.font.color
        if col is not None and col.rgb is not None:
            run.font.color.rgb = COLOR_MAP.get(str(col.rgb), col.rgb)
        if is_heading:
            run.font.color.rgb = ACCENT
        # underlines were the Gulim look; keep them only on real links
        if id(run._element) not in hl:
            run.underline = False

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style.element.rPr.rFonts.set(qn("w:ascii"), FONT)
style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

sec = doc.sections[0]
sec.left_margin = Cm(1.9)
sec.right_margin = Cm(1.7)
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.header_distance = Cm(0)
sec.footer_distance = Cm(0.9)

# --- full-bleed navy band, drawn as a page-anchored shape in the header ---
hdr = sec.header
hdr.is_linked_to_previous = False
hp = hdr.paragraphs[0]
hp.text = ""
hp.paragraph_format.space_after = Pt(0)
run = hp.add_run()
run.font.size = Pt(1)
# VML lives outside python-docx's namespace map, so build it from source
from docx.oxml import parse_xml

BAND_XML = (
    '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    '<v:rect style="position:absolute;margin-left:0;margin-top:0;width:595.5pt;'
    'height:12pt;z-index:-251658240;mso-position-horizontal:left;'
    'mso-position-horizontal-relative:page;mso-position-vertical:top;'
    'mso-position-vertical-relative:page" fillcolor="#0F3D57" stroked="f"/></w:pict>')
run._element.append(parse_xml(BAND_XML))

# --- footer: name on the left, page number on the right ---
ftr = sec.footer
ftr.is_linked_to_previous = False
fp = ftr.paragraphs[0]
fp.text = ""
fp.paragraph_format.space_before = Pt(0)
fp.paragraph_format.tab_stops.add_tab_stop(
    Cm(21 - 1.9 - 1.7), WD_TAB_ALIGNMENT.RIGHT)
r = fp.add_run("Seunghan Lee · Curriculum Vitae\t")
r.font.size = Pt(7.5)
r.font.color.rgb = RGBColor(0x8A, 0x8F, 0x96)
set_font(r)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
frun = OxmlElement("w:r")
rPr = OxmlElement("w:rPr")
szc = OxmlElement("w:sz")
szc.set(qn("w:val"), "15")
col = OxmlElement("w:color")
col.set(qn("w:val"), "8A8F96")
rf = OxmlElement("w:rFonts")
for a in ("w:ascii", "w:hAnsi"):
    rf.set(qn(a), FONT)
rPr.append(rf)
rPr.append(szc)
rPr.append(col)
frun.append(rPr)
txt = OxmlElement("w:t")
txt.text = "1"
frun.append(txt)
fld.append(frun)
fp._p.append(fld)

# --- lighter table rules, amber header row ---
for tbl in doc.tables:
    tblPr = tbl._tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    if tbl.style is not None and "Grid" in (tbl.style.name or ""):
        b = OxmlElement("w:tblBorders")
        for edge, sz, color in (("top", 10, "0F3D57"), ("bottom", 10, "0F3D57"),
                                ("left", 0, "FFFFFF"), ("right", 0, "FFFFFF"),
                                ("insideH", 4, "C9CFD6"), ("insideV", 0, "FFFFFF")):
            el = OxmlElement("w:" + edge)
            el.set(qn("w:val"), "single" if sz else "none")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            b.append(el)
        tblPr.append(b)
        for cell in tbl.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for shd in tcPr.findall(qn("w:shd")):
                tcPr.remove(shd)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "E6EDF2")
            tcPr.append(shd)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("styled", OUT, "headings:", n_head)


# ------------------------------------------------- embed the Lato family
def obfuscate(data, guid):
    """MS font obfuscation: XOR the first 32 bytes with the reversed GUID."""
    key = bytes.fromhex(guid.strip("{}").replace("-", ""))[::-1]
    out = bytearray(data)
    for i in range(32):
        out[i] ^= key[i % 16]
    return bytes(out)


FACES = [("embedRegular", "Lato-Regular.ttf"), ("embedBold", "Lato-Bold.ttf"),
         ("embedItalic", "Lato-Italic.ttf"), ("embedBoldItalic", "Lato-BoldItalic.ttf")]

tmp = OUT + ".tmp"
shutil.move(OUT, tmp)
zin = zipfile.ZipFile(tmp)
names = zin.namelist()
font_parts, rels, fonts_xml = [], [], []
for idx, (tag, fn) in enumerate(FACES, start=1):
    guid = "{%s}" % str(uuid.uuid4()).upper()
    blob = obfuscate(open(os.path.join(DL, fn), "rb").read(), guid)
    part = "word/fonts/font%d.odttf" % idx
    font_parts.append((part, blob))
    rid = "rIdFont%d" % idx
    rels.append('<Relationship Id="%s" Type="http://schemas.openxmlformats.org/'
                'officeDocument/2006/relationships/font" Target="fonts/font%d.odttf"/>'
                % (rid, idx))
    fonts_xml.append('<w:%s r:id="%s" w:fontKey="%s" w:subsetted="0"/>' % (tag, rid, guid))

font_table = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
    '<w:font w:name="Lato"><w:charset w:val="00"/><w:family w:val="swiss"/>'
    '<w:pitch w:val="variable"/>' + "".join(fonts_xml) + '</w:font>'
    '<w:font w:name="Calibri"><w:charset w:val="00"/><w:family w:val="swiss"/>'
    '<w:pitch w:val="variable"/></w:font></w:fonts>')

font_rels = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
             '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/'
             'relationships">' + "".join(rels) + '</Relationships>')

zout = zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED)
for item in zin.infolist():
    data = zin.read(item.filename)
    if item.filename == "[Content_Types].xml":
        s = data.decode("utf-8")
        if "odttf" not in s:
            s = s.replace("<Types ", '<Types ', 1)
            i = s.index(">", s.index("<Types")) + 1
            s = s[:i] + '<Default Extension="odttf" ContentType="application/vnd.openxml'\
                        'formats-officedocument.obfuscatedFont"/>' + s[i:]
        data = s.encode("utf-8")
    elif item.filename == "word/settings.xml":
        s = data.decode("utf-8")
        if "embedTrueTypeFonts" not in s:
            i = s.index(">", s.index("<w:settings")) + 1
            s = s[:i] + "<w:embedTrueTypeFonts/><w:saveSubsetFonts w:val=\"false\"/>" + s[i:]
        data = s.encode("utf-8")
    elif item.filename == "word/fontTable.xml":
        data = font_table.encode("utf-8")
    elif item.filename == "word/_rels/fontTable.xml.rels":
        data = font_rels.encode("utf-8")
    zout.writestr(item, data)

if "word/fontTable.xml" not in names:
    zout.writestr("word/fontTable.xml", font_table)
if "word/_rels/fontTable.xml.rels" not in names:
    zout.writestr("word/_rels/fontTable.xml.rels", font_rels)
for part, blob in font_parts:
    zout.writestr(part, blob)
zout.close()
zin.close()
os.remove(tmp)
print("embedded Lato:", ", ".join(f for _, f in FACES))
print("size", os.path.getsize(OUT), "bytes")
