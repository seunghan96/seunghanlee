# -*- coding: utf-8 -*-
"""Rebuild SeunghanLee_CV.pdf as an editable .docx (Gulim, same structure/colors)."""
import os
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/nfsdata/home/seunghan.lee/web/SeunghanLee_CV.docx"
PHOTO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "photo.png")

NAVY = RGBColor(0x1F, 0x49, 0x7D)
GREY = RGBColor(0x21, 0x21, 0x21)
GREY2 = RGBColor(0x33, 0x33, 0x33)
ORANGE = RGBColor(0xFF, 0x99, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)

doc = Document()

# ---------- page setup ----------
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(1.9)
sec.top_margin = sec.bottom_margin = Cm(2.2)

style = doc.styles["Normal"]
style.font.name = "Gulim"
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Gulim")
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15


def _fmt(run, bold=False, italic=False, underline=False, color=None, size=None, font="Gulim"):
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_hyperlink(par, text, url, size=None, color=BLUE, bold=False):
    part = par.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    link.append(new_run)
    par._p.append(link)
    run = docx.text.run.Run(new_run, par)
    run.text = text
    _fmt(run, bold=bold, underline=True, color=color, size=size)
    return run


def para(space_before=0, space_after=0, left=None, hanging=None, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if left is not None:
        pf.left_indent = Cm(left)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
    if align is not None:
        pf.alignment = align
    return p


def rich(p, parts, size=12):
    """parts: str | (text, opts-dict) | ('LINK', text, url, opts)"""
    for it in parts:
        if isinstance(it, str):
            _fmt(p.add_run(it), size=size)
        elif it[0] == "LINK":
            _, text, url = it[:3]
            o = it[3] if len(it) > 3 else {}
            add_hyperlink(p, text, url, size=o.get("size", size), bold=o.get("b", False))
        else:
            text, o = it
            r = _fmt(
                p.add_run(text),
                bold=o.get("b", False),
                italic=o.get("i", False),
                underline=o.get("u", False),
                color=o.get("c"),
                size=o.get("size", size),
            )
            if o.get("sup"):
                r.font.superscript = True
    return p


BULLETS = {1: "•\t", 2: "o\t", 3: "§\t"}
# measured from the original PDF: text at 36 pt from the margin, bullet 18 pt back
INDENT = {1: (1.27, 0.635), 2: (2.54, 0.635), 3: (3.81, 0.635)}
BULLET_PT = 10.0


def bullet(parts, level=1, size=12, space_before=0, space_after=0):
    left, hang = INDENT[level]
    p = para(space_before, space_after, left=left, hanging=hang)
    _fmt(p.add_run(BULLETS[level]), size=BULLET_PT)
    if isinstance(parts, str):
        parts = [parts]
    return rich(p, parts, size=size)


def add_bottom_border(p, sz=12, color="000000"):
    """Word draws a rule under each section heading."""
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    bdr.append(bottom)
    pPr.insert(0, bdr)


def h1(text, space_before=0, space_after=6, rule=True):
    p = para(space_before, space_after)
    p.paragraph_format.line_spacing = 1.0
    _fmt(p.add_run(text), bold=True, size=17)
    if rule:
        add_bottom_border(p)
    return p


def h2(text, space_before=8, space_after=2, size=12.8):
    p = para(space_before, space_after)
    _fmt(p.add_run(text), bold=True, underline=True, size=size)
    return p


def entry(parts, size=12.8, space_before=8, space_after=1, color=None, underline=True):
    """Sub-entry title line (numbered items, publication titles, ...)"""
    p = para(space_before, space_after)
    if isinstance(parts, str):
        parts = [(parts, {"b": True, "u": underline, "c": color, "size": size})]
    return rich(p, parts, size=size)


def blank(size=12):
    p = para()
    _fmt(p.add_run(""), size=size)
    return p


def page_break():
    doc.add_page_break()


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=12, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.paragraph_format.alignment = align
    _fmt(p.add_run(text), bold=bold, size=size, color=color)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


C = WD_ALIGN_PARAGRAPH.CENTER

# =====================================================================
# PAGE 1 — cover: photo + personal info, then careers summary
# =====================================================================
cover = doc.add_table(rows=1, cols=2)
cover.alignment = WD_TABLE_ALIGNMENT.LEFT
no_borders(cover)
cover.columns[0].width = Cm(4.5)
cover.columns[1].width = Cm(12.7)

pcell = cover.cell(0, 0)
pcell.width = Cm(4.5)
pp = pcell.paragraphs[0]
pp.paragraph_format.space_after = Pt(0)
pp.add_run().add_picture(PHOTO, width=Cm(3.97))

info = cover.cell(0, 1)
info.width = Cm(12.7)
info.text = ""
ip = info.paragraphs[0]
ip.paragraph_format.space_after = Pt(4)
_fmt(ip.add_run("Seunghan Lee"), bold=True, size=16)


INFO_INDENT = {1: (0.6, 0.6), 2: (1.7, 0.6)}


def info_bullet(parts, level=1, size=12):
    left, hang = INFO_INDENT[level]
    p = info.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(left)
    pf.first_line_indent = Cm(-hang)
    _fmt(p.add_run(BULLETS[level]), size=BULLET_PT)
    return rich(p, parts if not isinstance(parts, str) else [parts], size=size)


info_bullet("Date of Birth: July 29, 1996")
info_bullet("Country: Republic of Korea")
info_bullet("Military Service: Completed")
info_bullet("Contact Info.:")
p = info_bullet(["Email: "], level=2)
add_hyperlink(p, "seunghan9612@gmail.com", "mailto:seunghan9612@gmail.com", size=10)
p = info_bullet(["Linkedin: "], level=2)
add_hyperlink(
    p,
    "linkedin.com/in/seunghan-lee-0bba49169",
    "https://www.linkedin.com/in/seunghan-lee-0bba49169/",
    size=10,
)
p = info_bullet(["Homepage: "], level=2)
add_hyperlink(p, "seunghan96.github.io/seunghanlee", "https://seunghan96.github.io/seunghanlee/", size=10)
p = info_bullet(["Blog: "], level=2)
add_hyperlink(p, "seunghan96.github.io", "https://seunghan96.github.io/", size=10)
_fmt(p.add_run(" (1400+ AI-related posts)"), size=10)

blank()
blank()
h1("Careers (Brief)", space_after=10)

rich(para(space_before=4, space_after=2), [("LG AI Research", {"b": True, "u": True}), ": December 2025 – "])
bullet([("Role", {"b": True}), ": AI Scientist"])
bullet([("Division", {"b": True}), ": Data Intelligence Lab > Business Intelligence AI (BIAI)"])
rich(para(space_before=8, space_after=2), [("KRAFTON", {"b": True, "u": True}), ": May 2025 – December 2025"])
bullet([("Role", {"b": True}), ": Deep Learning Engineer/Researcher"])
bullet([("Division", {"b": True}), ": Deep Learning Division > AI Research > AI Agent Dept. > User Modeling AI Team"])
rich(para(space_before=8, space_after=2), [("Naver Cloud (internship)", {"b": True, "u": True}), ": March 2025 – May 2025"])
bullet([("Role", {"b": True}), ": AI Engineer/Researcher (Internship)"])
bullet([("Division", {"b": True}), ": HyperClova X"])
rich(para(space_before=8, space_after=2), [("SK Telecom (internship)", {"b": True, "u": True}), ": July 2019 – August 2019"])
bullet([("Role", {"b": True}), ": Data Scientist (Internship)"])
bullet([("Division", {"b": True}), ": ICT Infra Center > Data Science > Data Analytics (AI Home)"])
page_break()

# =====================================================================
# PAGE 2 — Education / Research Interests
# =====================================================================
h1("1. Education")
h2("Undergraduate", space_before=4)
rich(para(space_after=2), [("Yonsei University, Seoul", {"b": True}), " (March 2015 – August 2020)"])
bullet([("Majors", {"b": True}), ": Business Administration / Applied Statistics"])
bullet([("GPA (4.5 scale)", {"b": True}), ": Overall 4.19 (Business 4.31, Applied Statistics 4.35)"])
bullet([("Additional Information", {"b": True}), ": "])
bullet("Early graduation (7 semesters)", level=2)
bullet("Summa Cum Laude (Top 1% of the cohort)", level=2)

h2("Integrated Master's & PhD Program", space_before=12)
rich(para(space_after=2), [("Yonsei University, Seoul", {"b": True}), " (September 2020 – August 2025)"])
bullet([("Major", {"b": True}), ": Statistics and Data Science"])
bullet([("GPA (4.5 scale)", {"b": True}), ": 4.29"])
bullet([("Co-advisors", {"b": True}), ":"])
p = bullet(["Prof. Taeyoung Park ("], level=2)
add_hyperlink(p, "DSLab", "https://dslab-with.github.io/web/")
_fmt(p.add_run(")"), size=12)
p = bullet(["Prof. Kibok Lee ("], level=2)
add_hyperlink(p, "ML Lab", "https://ml.yonsei.ac.kr/")
_fmt(p.add_run(")"), size=12)

h1("2. Research Interests", space_before=18)
bullet([("Deep Learning for Time Series", {"b": True})])
bullet("Multimodal Time Series Forecasting", level=2)
bullet("Agentic Time Series Forecasting", level=2)
bullet("Time Series Foundation Models", level=2)
bullet("Financial Time Series", level=2)
bullet("Representation Learning", level=2)
bullet("Others (e.g., Self-Supervised Learning, Diffusion Models)", level=2)
bullet([("Deep Learning for Tabular / Graphs", {"b": True})])
page_break()

# =====================================================================
# PAGE 3-4 — Publications
# =====================================================================
h1("3. Publications", space_after=2)
p = para(space_after=6)
_fmt(p.add_run("("), size=10.5)
_fmt(p.add_run("*"), size=8).font.superscript = True
_fmt(p.add_run("Equally Contributed, "), size=10.5)
_fmt(p.add_run("†"), size=8).font.superscript = True
_fmt(p.add_run("Co-corresponding Authors)"), size=10.5)

ME = {"b": True, "u": True, "size": 11}
SUP = {"size": 7.5, "sup": True}
VEN = {"c": GREY, "size": 11, "i": True}


def pub(num, title, authors, venue_parts, extras=()):
    entry(str(num) + ". " + title, color=NAVY, underline=False)
    rich(para(space_after=0), authors, size=11)
    rich(para(space_after=0), venue_parts, size=11)
    for ex in extras:
        rich(para(space_after=0), ex, size=11)


ORAL = {"c": ORANGE, "size": 11}
LG_CO = ", Jun Seo, Jaehoon Lee, Sungdong Yoo, Minjae Kim, Tae Yoon Lim, Dongwan Kang, Hwanil Choi, SoonYoung Lee, Wonbin Ahn"
LG_CO2 = ", Jaehoon Lee, Jun Seo, Sungdong Yoo, Minjae Kim, Tae Yoon Lim, Dongwan Kang, Hwanil Choi, SoonYoung Lee, Wonbin Ahn"

pub(
    21,
    "FinVerse: Financial Time-Series Benchmark (2026)",
    ["- Jaehoon Lee, Jun Seo, ", ("Seunghan Lee", ME), ", Tae Yoon Lim, Dongwan Kang, Hwanil Choi, Minjae Kim, Sungdong Yoo, Junhyeok Kang, Sangjun Han, Soonyoung Lee, Wonbin Ahn"],
    ["- In ", ("arXiv preprint, 2026. [", VEN), ("LINK", "arxiv:2608.03259", "https://arxiv.org/abs/2608.03259", {"size": 11}), ("]", VEN)],
)
pub(
    20,
    "ReasonCast: Towards Explainable Time Series Forecasting with Reasoning (2026)",
    ["- ", ("Seunghan Lee", ME), ", Jun Seo, Jaehoon Lee, Junhyeok Kang, Sangjun Han, Sungdong Yoo, Minjae Kim, Tae Yoon Lim, Dongwan Kang, Hwanil Choi, Soonyoung Lee, Wonbin Ahn"],
    ["- In ", ("arXiv preprint, 2026. [", VEN), ("LINK", "arxiv:2608.01875", "https://arxiv.org/abs/2608.01875", {"size": 11}), ("]", VEN)],
)
pub(
    19,
    "Beyond Magnitude and Shape: A Direction-Aware Loss for Time Series Forecasting (2026)",
    ["- ", ("Seunghan Lee", ME), ", Jaehoon Lee, Jun Seo, Junhyeok Kang, Sangjun Han, Sungdong Yoo, Minjae Kim, Tae Yoon Lim, Dongwan Kang, Hwanil Choi, Soonyoung Lee, Wonbin Ahn"],
    ["- In ", ("arXiv preprint, 2026. [", VEN), ("LINK", "arxiv:2608.01857", "https://arxiv.org/abs/2608.01857", {"size": 11}), ("]", VEN)],
)
pub(
    18,
    "When Summaries Distort Decisions: Information Fidelity in LLM-Compressed Financial Analysis (2026)",
    ["- Hoyoung Lee, Suhwan Park, ", ("Seunghan Lee", ME), ", Jun Seo, Jaehoon Lee, Sungdong Yoo, Minjae Kim, CheolWon Na, Zhangyang Wang, Zach Golkhou, Minkyu Kim, Sotirios Sabanis, Alejandro Lopez-Lira, Dhagash Mehta, Soonyoung Lee, Chanyeol Choi, Wonbin Ahn, Yongjae Lee"],
    ["- In ", ("EMNLP, 2026. [", VEN), ("LINK", "arxiv:2606.29251", "https://arxiv.org/abs/2606.29251", {"size": 11}), ("]", VEN)],
)
pub(
    17,
    "AdaTKG: Adaptive Memory for Temporal Knowledge Graph Reasoning (2026)",
    ["- ", ("Seunghan Lee", ME), LG_CO],
    [
        "- In ",
        ("KDDW (GMLLM), 2026. ", VEN),
        ("Oral ", ORAL),
        ("[", VEN),
        ("LINK", "arxiv:2605.07121", "https://arxiv.org/abs/2605.07121", {"size": 11}),
        ("]", VEN),
    ],
)
pub(
    16,
    "FinSTaR: Towards Financial Reasoning with Time Series Reasoning Models (2026)",
    ["- ", ("Seunghan Lee", ME), LG_CO],
    [
        "- In ",
        ("EMNLP, 2026. [", VEN),
        ("LINK", "arxiv:2605.03460", "https://arxiv.org/abs/2605.03460", {"size": 11}),
        ("]", VEN),
    ],
    extras=[
        [
            ("- Preliminary version: KDDW, 2026. ", VEN),
            ("Oral", ORAL),
        ]
    ],
)
pub(
    15,
    "Not All Retrievals are Useful: Cross-Attention for Input-Aware RAG in Time Series Forecasting (2026)",
    ["- ", ("Seunghan Lee", ME), LG_CO2],
    [
        "- In ",
        ("KDDW (MILETS), 2026. ", VEN),
        ("Oral ", ORAL),
        ("[", VEN),
        ("LINK", "arxiv:2603.14709", "https://arxiv.org/abs/2603.14709", {"size": 11}),
        ("]", VEN),
    ],
)
pub(
    14,
    "Rethinking Multimodal Fusion for Time Series: Auxiliary Modalities Need Constrained Fusion (2026)",
    ["- ", ("Seunghan Lee", ME), LG_CO],
    [
        "- In ",
        ("KDDW (MILETS), 2026. ", VEN),
        ("Oral ", ORAL),
        ("[", VEN),
        ("LINK", "arxiv:2603.22372", "https://arxiv.org/abs/2603.22372", {"size": 11}),
        ("]", VEN),
    ],
)
pub(
    13,
    "FinTexTS: Financial Text-Paired Time-Series Dataset via Semantic-Based and Multi-Level Pairing (2026)",
    [
        "- Jaehoon Lee, Suhwan Park, Tae Yoon Lim, ",
        ("Seunghan Lee", ME),
        ", Jun Seo, Dongwan Kang, Hwanil Choi, Minjae Kim, Sungdong Yoo, SoonYoung Lee, Yongjae Lee, Wonbin Ahn",
    ],
    ["- In ", ("KDD, 2026. [", VEN), ("LINK", "arxiv:2603.02702", "https://arxiv.org/abs/2603.02702", {"size": 11}), ("]", VEN)],
)
pub(
    12,
    "Mitigating Label Shift in Tabular In-Context Learning via Test-Time Posterior Adjustment (2026)",
    ["- ", ("Seunghan Lee", ME), LG_CO2],
    ["- In ", ("ICML, 2026. [", VEN), ("LINK", "arxiv:2605.04363", "https://arxiv.org/abs/2605.04363", {"size": 11}), ("]", VEN)],
)
pub(
    11,
    "Dataset-Driven Channel Masks in Transformers for Multivariate Time Series (2026)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park, Kibok Lee"],
    ["- In ", ("ICASSP, 2026. [", VEN), ("LINK", "arxiv:2410.23222", "https://arxiv.org/abs/2410.23222", {"size": 11}), ("]", VEN)],
)
pub(
    10,
    "Soft Contrastive Learning for Irregular Time Series (2025)",
    ["- Junghoon Lim, ", ("Seunghan Lee", ME), ", Taeyoung Park"],
    ["- In ", ("ICMLW (Foundation Models for Structured Data), 2025.", VEN)],
)
pub(
    9,
    "Channel Normalization for Time Series Channel Identification (2025)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park", ("†", SUP), ", Kibok Lee", ("†", SUP)],
    ["- In ", ("ICML, 2025. [", VEN), ("LINK", "arxiv:2506.00432", "https://arxiv.org/abs/2506.00432", {"size": 11}), ("]", VEN)],
)
pub(
    8,
    "Sequential Order-Robust Mamba for Time Series Forecasting (2024)",
    ["- ", ("Seunghan Lee", ME), ("*", SUP), ", Juri Hong", ("*", SUP), ", Kibok Lee", ("†", SUP), ", Taeyoung Park", ("†", SUP)],
    ["- In ", ("NeurIPSW (Time Series in the Age of Large Models) , 2024. [", VEN), ("LINK", "arxiv:2410.23356", "https://arxiv.org/abs/2410.23356", {"size": 11}), ("]", VEN)],
)
pub(
    7,
    "Partial Channel Dependence with Channel Masks for Time Series Foundation Models (2024)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park", ("†", SUP), ", Kibok Lee", ("†", SUP)],
    [
        "- In ",
        ("NeurIPSW (Time Series in the Age of Large Models), 2024. ", VEN),
        ("Oral Presentation (Top 5) ", {"c": ORANGE, "size": 11}),
        ("[", VEN),
        ("LINK", "arxiv:2410.23222", "https://arxiv.org/abs/2410.23222", {"size": 11}),
        ("]", VEN),
    ],
)
pub(
    6,
    "Adaptive Noise Schedule for Time Series Diffusion Models (2024)",
    ["- ", ("Seunghan Lee", ME), ", Kibok Lee", ("†", SUP), ", Taeyoung Park", ("†", SUP)],
    ["- In ", ("NeurIPS, 2024. [", VEN), ("LINK", "arxiv:2410.14488", "https://arxiv.org/abs/2410.14488", {"size": 11}), ("]", VEN)],
    extras=[["- ", ("Outstanding Paper Award", {"c": ORANGE, "size": 11}), " from the ", ("Journal of Korean Artificial Intelligence Association", {"b": True, "size": 11})]],
)
pub(
    5,
    "Learning to Embed Time Series Patches Independently (2024)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park, Kibok Lee"],
    ["- In ", ("ICLR, 2024. [", VEN), ("LINK", "arxiv:2312.16427", "https://arxiv.org/abs/2312.16427", {"size": 11}), ("]", VEN)],
    extras=[
        [
            ("- Preliminary version: NeurIPSW (Self-Supervised Learning: Theory and Practice), 2023. ", VEN),
            ("Oral Presentation (Top 4)", {"c": ORANGE, "size": 11}),
        ]
    ],
)
pub(
    4,
    "Soft Contrastive Learning for Time Series (2024)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park, Kibok Lee"],
    [
        "- In ",
        ("ICLR, 2024. ", VEN),
        ("Spotlight (366/7262=5%) ", {"c": ORANGE, "size": 11}),
        ("[", VEN),
        ("LINK", "arxiv:2312.16424", "https://arxiv.org/abs/2312.16424", {"size": 11}),
        ("]", VEN),
    ],
    extras=[[("- Preliminary version: NeurIPSW (Self-Supervised Learning: Theory and Practice), 2023.", VEN)]],
)
pub(
    3,
    "Hierarchical Multi-Task Learning with Self-Supervised Auxiliary Task (2024)",
    ["- ", ("Seunghan Lee", ME), ", Taeyoung Park"],
    ["- In ", ("The Korean Journal of Applied Statistics 37(5): 631-641, 2024.", VEN)],
)
pub(
    2,
    "MAD-GL2: Multimodal Adaptive Dynamic Graph Learning with Global and Local Features for Multivariate Time Series Forecasting (2025+)",
    ["- ", ("Seunghan Lee", ME), ("*", SUP), ", Kibok Lee", ("*", SUP), ", Taeyoung Park"],
    ["- Under review"],
)
pub(
    1,
    "Improving Gibbs Sampler (2022)",
    ["- Taeyoung Park, ", ("Seunghan Lee", ME)],
    ["- In ", ("Wiley Interdisciplinary Reviews: Computational Statistics 14(2): e1546, 2022.", {"c": RGBColor(0x22, 0x22, 0x22), "size": 11})],
    extras=[["- Number of Citations (Google Scholar, 2025.02.16): 11"]],
)
page_break()

# =====================================================================
# PAGE 5 — Careers (Details)
# =====================================================================
h1("4. Careers (Details)", space_after=6)

entry("1. LG AI Research", space_before=2)
bullet([("Period", {"b": True}), ": December 2025 – "])
bullet([("Role", {"b": True}), ": AI Scientist"])
bullet([("Division", {"b": True}), ": Data Intelligence Lab > Business Intelligence AI (BIAI)"])
bullet([("Tasks", {"b": True}), ": "])
bullet("① TS foundation model for financial assets", level=2)
bullet("② Financial TS forecasting (Equity, Macro, Fundamental, etc.)", level=2)
bullet("③ Multimodal TS forecasting (TS + Text)", level=2)
bullet("④ Agentic TS forecasting", level=2)
bullet("⑤ Supply Chain Knowledge Graph", level=2)
bullet([("Key Algorithms", {"b": True}), ": TS Foundation Models, Multimodal Learning, LLM Agents, Knowledge Graph"])
bullet([("Programming Language", {"b": True}), ": Python"])

entry("2. KRAFTON")
bullet([("Period", {"b": True}), ": May 2025 – December 2025"])
bullet([("Role", {"b": True}), ": Deep Learning Engineer/Researcher"])
bullet([("Division", {"b": True}), ": Deep Learning Division > AI Research > AI Agent Dept. > User Modeling AI Team"])
bullet([("Tasks", {"b": True}), ": "])
bullet("① PUBG anti-cheat modeling with tabular and time-series DL", level=2)
bullet("② PUBG Global Championship (PGC) win-rate prediction with tabular DL and survival analysis", level=2)
bullet([("Key Algorithms", {"b": True}), ": Tabular DL, Time Series DL, Survival Analysis"])
bullet([("Programming Language", {"b": True}), ": Python"])

entry("3. Naver Cloud (internship)")
bullet([("Period", {"b": True}), ": March 2025 – May 2025"])
bullet([("Division", {"b": True}), ": HyperClova X"])
bullet([("Tasks", {"b": True}), ": Multi-modal Pretraining."])
bullet("① Model training tasks for adding vision capabilities to a multimodal backbone", level=2)
bullet("② Exploration of the impact of various multimodal backbones on performance", level=2)
bullet("③ Efficient pretraining recipe exploration", level=2)
bullet([("Key Algorithms", {"b": True}), ": LLaVA, Qwen, DeepSeek"])
bullet([("Programming Language", {"b": True}), ": Python"])

entry("4. SK Telecom (internship)")
bullet([("Period", {"b": True}), ": July 2019 – August 2019"])
bullet([("Division", {"b": True}), ": ICT Infra Center > Data Science > Data Analytics (AI Home)"])
bullet([("Tasks", {"b": True}), ": Data analysis and service planning for SKT’s air quality info service."])
bullet("① Analyzed air quality data and addressed measurement errors.", level=2)
bullet("② Designed strategies based on data analysis of parents in their 30s.", level=2)
bullet("③ Planned air quality reports, including data visualization.", level=2)
bullet([("Key Algorithms", {"b": True}), ": Clustering, DNN"])
bullet([("Programming Language", {"b": True}), ": Python"])

entry("5. Yonsei Univ. Department of Computational Science & Engineering")
bullet([("Period", {"b": True}), ": December 2019 – February 2020"])
bullet([("Division", {"b": True}), ": Data Science & Deep Learning > Network Embedding"])
bullet([("Tasks", {"b": True}), ": Implemented research papers and conducted related seminars."])
bullet([("Key Algorithms", {"b": True}), ": Graph Neural Network"])
bullet([("Programming Language", {"b": True}), ": Python"])
page_break()

# =====================================================================
# PAGE 6 — Industry-Academia / Invited Talks / IP
# =====================================================================
h1("5. Industry-Academia Collaboration", space_after=10)

entry("1. Amore Pacific", space_before=2)
bullet([("Period", {"b": True}), ": February 2020 – December 2022"])
bullet([("[2020 Project]", {"b": True}), " User and purchase behavior analysis for Amore Pacific."])
bullet("① Analyzed purchasing behavior and conducted customer segmentation.", level=2)
bullet("② Predicted sales of products.", level=2)
bullet("③ Analyzed web log data to understand online user behavior patterns.", level=2)
bullet([("Key Algorithms", {"b": True}), ": Random Forest, Logistic Regression, Hierarchical Clustering"], level=2)
bullet([("[2021–2022 Project]", {"b": True}), " Development of Context Brand Score (CBS)."])
bullet("Developed a brand index that considers consumer sentiment on cosmetic product attributes using text data (social buzz, reviews).", level=2)
bullet([("Key Algorithms", {"b": True}), ": Hierarchical BERT, Aspect-Based Sentiment Analysis."], level=2)
bullet([("Programming Language", {"b": True}), ": Python, R"])

entry("2. Hyundai Mobis")
bullet([("Period", {"b": True}), ": June 2022 – December 2022"])
bullet([("Tasks", {"b": True}), ": "])
bullet("① Short- and long-term demand forecasting for low-circulation and eco-friendly parts of vehicles (Zero-inflated TS).", level=2)
bullet("② Utilizing textual data and tabular data for TS forecasting.", level=2)
bullet([("Key Algorithms", {"b": True}), ": GNN, GAN, Domain Adaptation, SSL with Tabular Data"])
bullet([("Programming Language", {"b": True}), ": Python"])

h1("6. Academic Service", space_before=20, space_after=6)
bullet([("Reviewer", {"b": True}), ": ICLR (2025, 2026), ICML (2024, 2025, 2026), NeurIPS (2024, 2025), EMNLP (2026), TMLR (2026)"])

h1("7. Invited Talks", space_before=20, space_after=6)
bullet("2024.08. Korean Artificial Intelligence Association (2024 Summer Conference)")

h1("8. Intellectual Property", space_before=20, space_after=6)
bullet("2023.10: Review evaluation method using review quality metrics")
bullet("2024.12: Universal time series analysis method and system for irregular multivariate time series")
page_break()

# =====================================================================
# PAGE 7 — Programming Languages / Instructor & TA
# =====================================================================
h1("9. Programming Languages")
bullet([("Proficient", {"b": True}), ": Python (Pytorch, Tensorflow, Numpy, Sklearn)"])
bullet([("Intermediate", {"b": True}), ": SQL, R"])
bullet([("Basic", {"b": True}), ": Docker, Kubernetes"])

h1("10. Instructor, Teaching Assistant (TA)", space_before=20, space_after=6)
entry("1. [Instructor] Yonsei University Atmospheric Science Department – Data Analysis")
bullet([("Period", {"b": True}), ": November 2020 – February 2021"])
bullet([("Tasks", {"b": True}), ": Provided statistical and clustering analysis training for graduate students and assisted with data analysis."])

entry("2. [Instructor] KB Bank – Computer Vision / GAN")
bullet([("Period", {"b": True}), ": July 2021"])
bullet([("Tasks", {"b": True}), ": Provided a lecture on computer vision and generative models"])

entry("3. [Instructor] Hanwha Ocean – Data Analysis & Statistics")
bullet([("Period", {"b": True}), ": July 2023"])
bullet([("Tasks", {"b": True}), ": Provided a lecture on below contents regarding statistical analysis"])

entry("4. [TA] Python Programming & Web Crawling")
bullet([("Period", {"b": True}), ": March, September 2021,2022,2023,2024"])
bullet([("Tasks", {"b": True}), ": Teaching assistant for Python programming and web crawling"])

entry("5. [TA] SK Hynix – Bayesian Statistics & ML Course")
bullet([("Period", {"b": True}), ": April 2021 – June 2021"])
bullet([("Tasks", {"b": True}), ": Conducted coding exercises for Bayesian Statistics and ML."])
page_break()

# =====================================================================
# PAGE 8-9 — Academic Society & Club
# =====================================================================
h1("11. Academic Society & Club", space_after=10)

entry("1. ISSU (Yonsei University IT Management Strategy Society)", space_before=2)
bullet([("Period", {"b": True}), ": March 2018 – December 2018"])
bullet([("Regular Sessions", {"b": True}), ": Conducted industry-academia collaboration projects with IT companies, including Naver Papago, Naver VIBE, Send Anywhere, and Amanda."])
bullet([("Study Sessions", {"b": True}), ":"])
p = bullet(["1st Semester: Contributed an article on AI in the Arts to IT Chosun ("], level=2)
add_hyperlink(p, "Link", "http://it.chosun.com/site/data/html_dir/2018/06/28/2018062802823.html")
_fmt(p.add_run(")."), size=12)
bullet("2nd Semester: Served as a Python mentor.", level=2)

entry("2. BITAMIN (Big Data Analytics Club)")
bullet([("Period", {"b": True}), ": August 2018 – March 2019"])
bullet([("Regular Sessions", {"b": True}), ": Focused on machine learning and text mining."])
bullet([("Study Sessions", {"b": True}), ": Conducted three machine learning projects."])

entry("3. Data Science Lab (Yonsei University Data Science Society)")
bullet([("Period", {"b": True}), ": February 2019 – June 2020"])
bullet([("The 1st President", {"b": True}), ": Organized regular sessions and study groups on ML and DL."])
bullet([("Regular Sessions", {"b": True}), ": "])
bullet("① Studied the entire data science workflow, from data collection and preprocessing to modeling.", level=2)
bullet("② Conducted training sessions for society members on:", level=2)
bullet("Data Preprocessing, Association Analysis, SVM, Boosting, CNN", level=3)
bullet([("Study Sessions", {"b": True}), ": Group projects and data science competitions."])

blank()
club = doc.add_table(rows=4, cols=4, style="Table Grid")
club.alignment = WD_TABLE_ALIGNMENT.CENTER
club.autofit = False
widths = [Cm(3.2), Cm(6.4), Cm(3.2), Cm(4.4)]
for i, w in enumerate(widths):
    club.columns[i].width = w
for row in club.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w
hdr = ["Name", "Details", "Period", "Etc."]
for i, t in enumerate(hdr):
    cell_text(club.cell(0, i), t, bold=True, size=12.8, align=C)
    shade(club.cell(0, i), "D9E2F3")
rows = [
    ("ISSU", "Yonsei Univ. IT Management Strategy Society", "2018.03 - 2018.12", "Industry-academia collaboration projects.\nPython mentor."),
    ("Bitamin", "Big Data Analytics Club", "2018.07 - 2019.06", ""),
    ("Data Science Lab", "Yonsei Univ. Data Science Society", "2019.01 - 2020.06", "The 1st President"),
]
for r, vals in enumerate(rows, start=1):
    for c, v in enumerate(vals):
        if "\n" in v:
            cell_text(club.cell(r, c), v.split("\n")[0])
            pp2 = club.cell(r, c).add_paragraph()
            pp2.paragraph_format.space_after = Pt(0)
            pp2.paragraph_format.line_spacing = 1.0
            _fmt(pp2.add_run(v.split("\n")[1]), size=12)
        else:
            cell_text(club.cell(r, c), v, align=C if c in (0, 2) else None)

entry("4. OWOP (One Week One Paper)", space_before=16)
bullet([("Period", {"b": True}), ": January 2021 – June 2021"])
bullet([("Tasks", {"b": True}), ": Conducted a research paper study group with five graduate students from the Department of Statistics and Data Science at Yonsei University."])
bullet([("Topics", {"b": True}), ": AI and statistics-related subjects, including CV, NLP, BNN, Recommender Systems, Deep Generative Models (VAE, Normalizing Flow)."])

entry("5. Deep Learning Paper Reading Study")
bullet([("Period", {"b": True}), ": May 2021 – September 2021"])
bullet([("Tasks", {"b": True}), ": Conducted a research paper study group with five graduate students from the Department of Statistics and Data Science at Yonsei University."])
bullet([("Topics", {"b": True}), ": Meta Learning, Continual Learning, Interpretable/Reliable Learning"])
page_break()

# =====================================================================
# PAGE 10 — Certifications & Awards
# =====================================================================
h1("12. Certifications & Awards", space_after=10)
h2("Certifications", space_before=2, space_after=4)

cert_rows = [
    ("CS/Data", "ADSP (Associate Data Science Professional)", "-", "2018.09"),
    ("", "Computer Proficiency Level 1", "", "2019.03"),
    ("", "SQLD (SQL Developer)", "", "2020.06"),
    ("", "Big Data Analyst", "", "2021.07"),
    ("", "AWS Cloud Practitioner", "", "2021.08"),
    ("Linguistics", "TOEIC", "965/990", "2020.08"),
    ("Etc.", "Korean History Proficiency Test", "Level 1", "2012.08"),
    ("", "TESAT (Korean Economic Newspaper)", "S-Level", "2013.05"),
]
cert = doc.add_table(rows=1 + len(cert_rows), cols=4, style="Table Grid")
cert.autofit = False
cw = [Cm(2.2), Cm(9.2), Cm(2.6), Cm(3.2)]
for i, w in enumerate(cw):
    cert.columns[i].width = w
for row in cert.rows:
    for i, w in enumerate(cw):
        row.cells[i].width = w
for i, t in enumerate(["Field", "Certifications", "Score", "Date"]):
    cell_text(cert.cell(0, i), t, bold=True, size=12.8, align=C)
    shade(cert.cell(0, i), "D9E2F3")
for r, vals in enumerate(cert_rows, start=1):
    for c, v in enumerate(vals):
        cell_text(cert.cell(r, c), v, align=C if c != 1 else None)
cert.cell(1, 0).merge(cert.cell(5, 0))
cell_text(cert.cell(1, 0), "CS/Data", align=C)
cert.cell(1, 2).merge(cert.cell(5, 2))
cell_text(cert.cell(1, 2), "-", align=C)
cert.cell(7, 0).merge(cert.cell(8, 0))
cell_text(cert.cell(7, 0), "Etc.", align=C)

h2("Awards", space_before=16, space_after=4)
award_rows = [
    ("Academic", "Academic Excellence Awards", "", "2018.12 / 2019.06"),
    ("", "Summa Cum Laude Graduation", "", "2020.08"),
    ("CS", "Yonsei Data Science Competition", "2nd", "2019.11"),
    ("", "Sony Pictures Audience Prediction", "4th", "2020.01"),
    ("", "Yonsei Data Science Competition", "2nd", "2021.12"),
    ("", "BK Winter Academic Conference", "4th", "2022.12"),
    ("", "Yonsei Data Science Competition", "1st", "2023.02"),
    ("Economy", "TESAT Team Excellence Award", "2nd", "2013.05"),
]
aw = doc.add_table(rows=1 + len(award_rows), cols=4, style="Table Grid")
aw.autofit = False
for i, w in enumerate(cw):
    aw.columns[i].width = w
for row in aw.rows:
    for i, w in enumerate(cw):
        row.cells[i].width = w
for i, t in enumerate(["Field", "Awards", "Rank", "Date"]):
    cell_text(aw.cell(0, i), t, bold=True, size=12.8, align=C)
    shade(aw.cell(0, i), "D9E2F3")
for r, vals in enumerate(award_rows, start=1):
    for c, v in enumerate(vals):
        cell_text(aw.cell(r, c), v, align=C if c != 1 else None)
aw.cell(1, 0).merge(aw.cell(2, 0))
cell_text(aw.cell(1, 0), "Academic", align=C)
aw.cell(1, 2).merge(aw.cell(2, 2))
cell_text(aw.cell(1, 2), "", align=C)
aw.cell(3, 0).merge(aw.cell(7, 0))
cell_text(aw.cell(3, 0), "CS", align=C)
page_break()

# =====================================================================
# PAGE 11 — Projects
# =====================================================================
h1("13. Projects", space_after=8)
entry("1. Fine-tuning LLM", space_before=2)
bullet([("Period", {"b": True}), ": February 2025 – March 2025"])
bullet([("Tasks", {"b": True}), ": "])
bullet("(1) Fine-tuning LLaMA with Q-LoRA (Single-GPU)", level=2)
bullet("(2) Fine-tuning LLaMA with Q-LoRA & FFT (Multi-GPU)", level=2)
bullet("Axolotl, code-based", level=3)
bullet("Distributed training: MP, FSDP, ZeRO", level=3)
bullet("(3) Fine-tuning LLaMA for the insurance domain", level=2)
bullet("One-cycle project: Data construction – SFT – DPO – Inference", level=3)
bullet("Extract text from PDFs using OCR, generate additional questions with WizardLM (i.e., evolving), and build SFT & DPO datasets", level=3)

entry([("2. Hands-on practice with Ollama, RAG pipeline, Vector DB, and UI/UX tools", {"b": True, "c": GREY2, "size": 12.8})])
bullet([("Period", {"b": True}), ": February 2025 – March 2025"])
bullet([("Tasks", {"b": True}), ": "])
bullet([("(1) Ollama + RAG (PDF upload) + ChromaDB + Streamlit", {"c": GREY2})], level=2)
bullet([("(2) Ollama + RAG (Repository) + ChromaDB + Chainlit", {"c": GREY2})], level=2)
bullet([("(3) Ollama + RAG (Repository) + ChromaDB + Gradio", {"c": GREY2})], level=2)
page_break()

# =====================================================================
# PAGE 12-14 — Competitions
# =====================================================================
h1("14. Competitions", space_after=8)


def comp(title, period, tasks, subs=(), algos=None, result=None, algo_label="Key Algorithms"):
    entry(title)
    bullet([("Period", {"b": True}), ": " + period])
    bullet([("Tasks", {"b": True}), ": " + tasks])
    for s in subs:
        bullet(s, level=2)
    if algos:
        bullet([(algo_label, {"b": True}), ": " + algos])
    if result:
        bullet([("Result", {"b": True}), ": " + result])


comp(
    "1. [Dacon] 2019 KBO Baseball Player Performance Prediction",
    "January 2019 – February 2019",
    "Predicting KBO baseball player performance for the first half of 2019.",
    algos="MA, ARIMA, Random Forest",
)
comp(
    "2. [Yonsei Data Science Competition] Embrain",
    "March 2019 – May 2019",
    "Developing a marketing strategy for highly meticulous consumers.",
    algos="Random Forest, NN",
    result="4th",
)
comp(
    "3. [Yonsei Data Science Competition] PACKUS",
    "September 2019 – November 2019",
    "",
    subs=[
        "① Customer segmentation and personalized strategy development.",
        "② Future sales prediction for key products.",
        "③ Created a dashboard for CRM using R Shiny.",
    ],
    algos="RNN, XGBoost, LightGBM, SOM, Association Analysis",
    result="2nd",
)
comp(
    "4. [Sony Pictures] Movie Audience Prediction",
    "November 2019 – January 2020",
    "Predicting cumulative audience numbers for the first 14 days after a movie's release.",
    algos="Random Forest, XGBoost, Stacking, NN",
    result="4th",
)
comp(
    "5. [Kakao] Melon Playlist Continuation",
    "May 2019 – July 2019",
    "Predicting missing songs and tags in a given playlist.",
    algos="Collaborative Filtering, Spectrogram Analysis",
)
comp(
    "6. [IGA Works] CTR Prediction",
    "December 2019 – February 2020",
    "Predicting click-through probability when a user is exposed to an ad.",
    algos="Deep CTR",
    algo_label="Key Algorithm",
)
comp(
    "7. [Dacon] AI Competition for Psychological Tendency Prediction",
    "September 2020 – November 2020",
    "Developing an algorithm for psychological test analysis and voter prediction.",
    algos="AutoML, Deep CTR",
)
comp(
    "8. [BigCon Test] NS Home Shopping Schedule Optimization",
    "July 2020 – October 2020",
    "Optimizing home shopping schedules to maximize sales.",
    algos="CatBoost, LightGBM, Bayesian Optimization",
)
comp(
    "9. [Hyundai Industries] Big Data/AI Competition",
    "January 2021 – February 2021",
    "Predicting manufacturing process task durations and optimizing task allocation.",
    algos="Bayesian Neural Network, LightGBM",
    result="Finalist (18th place out of 284 teams)",
)
comp(
    "10. [Yonsei Data Science Competition] Hyodol",
    "September 2021 – December 2021",
    "Customer clustering and predicting optimal engagement times for personalized care.",
    algos="Multi-task Learning, Self-Organizing Map",
    result="2nd",
)
comp(
    "11. [Conference] 2022 Winter BK Academic Conference",
    "December 2022",
    "Multivariate time series forecasting using Spatio-Temporal GNN.",
    algos="Spatio-Temporal GNN",
    result="4th",
)
comp(
    "12. [Yonsei Data Science Competition] KCB",
    "December 2022 – February 2023",
    "Multi-task learning with self-supervised learning.",
    subs=[
        "① Regression: Credit score prediction",
        "② Classification: Credit rating prediction",
        "③ Time Series Forecasting: Future revenue prediction",
        "④ Clustering: Customer segmentation",
    ],
    algos="SSL with Tabular Data",
    result="1st",
)

# fix: item 3 has no inline task text
doc.save(OUT)
print("saved", OUT)
